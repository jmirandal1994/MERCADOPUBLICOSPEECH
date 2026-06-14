#!/usr/bin/env python3
"""
Preparador de Anexos - Speech Psychology SPA
Llena documentos Word/PDF con datos de la empresa e inserta firma+timbre
"""

import os
import sys
import shutil
from pathlib import Path
from datetime import datetime
import anthropic
import base64

# Datos empresa
EMPRESA = {
    "razon_social": "SPEECH PSYCHOLOGY SPA",
    "rut": "78.254.509-4",
    "rut_sin_puntos": "78254509-4",
    "giro": "Servicios de Atención de Salud Humana Médicos y No Médicos",
    "codigo_actividad": "869091",
    "direccion": "Vic. Mackenna Pte. 6843 Of. 504 Torre A",
    "depto": "504",
    "bloque": "Torre A",
    "comuna": "La Florida",
    "region": "Metropolitana de Santiago",
    "representante_legal": "Jorge Ignacio Miranda López",
    "rut_rep_legal": "18.953.982-7",
    "cargo_rep_legal": "Representante Legal",
    "profesion": "Fonoaudiólogo",
    "email": "contacto@speechpsychology.cl",
    "telefono": "+56978426971",
    "tipo_empresa": "Sociedad por Acciones (SpA)",
    "segmento": "Micro Empresa",
    "fecha_inicio": "22/09/2025",
    "categoria_tributaria": "Primera Categoría",
    "afecto_iva": "No",
}

ANTHROPIC_API_KEY = "C1148555-988E-40E5-9115-809B36F23168"

# Ruta a la imagen de firma y timbre
FIRMA_TIMBRE_PATH = Path(__file__).parent.parent / "firma_timbre.png"


def analizar_anexo_con_ia(archivo_path):
    """Usa Claude para analizar el contenido del anexo y extraer campos a llenar."""
    
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    
    # Leer el archivo
    with open(archivo_path, "rb") as f:
        contenido = f.read()
    
    ext = Path(archivo_path).suffix.lower()
    
    if ext == ".pdf":
        # Enviar PDF a Claude
        b64 = base64.standard_b64encode(contenido).decode()
        
        mensaje = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": b64
                        }
                    },
                    {
                        "type": "text",
                        "text": f"""Analiza este documento (anexo de licitación pública chilena).

Empresa postulante:
{chr(10).join([f'- {k}: {v}' for k, v in EMPRESA.items()])}

Identifica:
1. Todos los campos que hay que llenar (nombre del campo y valor a poner)
2. Si hay que poner firma y/o timbre (dónde)
3. Requisitos especiales o declaraciones juradas
4. Documentos adicionales que menciona

Responde en formato JSON:
{{
  "titulo_documento": "nombre del anexo",
  "campos_a_llenar": [
    {{"campo": "nombre del campo", "valor": "valor a ingresar", "ubicacion": "descripcion de donde esta"}}
  ],
  "requiere_firma": true/false,
  "requiere_timbre": true/false,
  "declaraciones": ["declaracion 1", "declaracion 2"],
  "documentos_adicionales": ["doc 1", "doc 2"],
  "instrucciones_especiales": "notas importantes"
}}"""
                    }
                ]
            }]
        )
    else:
        # Texto plano
        try:
            texto = contenido.decode("utf-8", errors="ignore")
        except:
            texto = str(contenido[:3000])
        
        mensaje = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2000,
            messages=[{
                "role": "user",
                "content": f"""Analiza este documento (anexo de licitación):

{texto[:3000]}

Empresa postulante:
{chr(10).join([f'- {k}: {v}' for k, v in EMPRESA.items()])}

Responde en JSON con los campos a llenar y requisitos."""
            }]
        )
    
    texto_resp = mensaje.content[0].text.strip()
    
    # Limpiar markdown
    if "```" in texto_resp:
        partes = texto_resp.split("```")
        for p in partes:
            if p.startswith("json"):
                texto_resp = p[4:].strip()
                break
            elif p.strip().startswith("{"):
                texto_resp = p.strip()
                break
    
    import json
    try:
        return json.loads(texto_resp)
    except:
        return {"error": "No se pudo analizar", "texto_raw": texto_resp}


def llenar_anexo_word(archivo_path, analisis):
    """Llena un documento Word con los datos de la empresa."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        import copy
        
        doc = Document(archivo_path)
        campos = analisis.get("campos_a_llenar", [])
        
        reemplazos = {}
        for campo in campos:
            nombre = campo.get("campo", "").lower()
            valor = campo.get("valor", "")
            
            # Mapeo de campos comunes
            mapeo = {
                "razón social": EMPRESA["razon_social"],
                "razon social": EMPRESA["razon_social"],
                "nombre empresa": EMPRESA["razon_social"],
                "rut empresa": EMPRESA["rut"],
                "rut": EMPRESA["rut"],
                "giro": EMPRESA["giro"],
                "dirección": EMPRESA["direccion"],
                "direccion": EMPRESA["direccion"],
                "comuna": EMPRESA["comuna"],
                "región": EMPRESA["region"],
                "representante legal": EMPRESA["representante_legal"],
                "nombre representante": EMPRESA["representante_legal"],
                "rut representante": EMPRESA["rut_rep_legal"],
                "cargo": EMPRESA["cargo_rep_legal"],
                "correo": EMPRESA["email"],
                "email": EMPRESA["email"],
                "teléfono": EMPRESA["telefono"],
                "telefono": EMPRESA["telefono"],
                "fecha": datetime.now().strftime("%d/%m/%Y"),
            }
            
            for k, v in mapeo.items():
                if k in nombre:
                    reemplazos[campo.get("campo", "")] = v
                    break
            else:
                reemplazos[campo.get("campo", "")] = valor
        
        # Buscar y reemplazar en párrafos
        for para in doc.paragraphs:
            for campo, valor in reemplazos.items():
                if campo and campo in para.text:
                    # Reemplazar manteniendo formato
                    for run in para.runs:
                        if campo in run.text:
                            run.text = run.text.replace(campo, str(valor))
        
        # Buscar en tablas
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for campo, valor in reemplazos.items():
                        if campo and campo in celda.text:
                            for para in celda.paragraphs:
                                for run in para.runs:
                                    if campo in run.text:
                                        run.text = run.text.replace(campo, str(valor))
        
        # Insertar firma e imagen si se requiere
        if analisis.get("requiere_firma") and FIRMA_TIMBRE_PATH.exists():
            # Agregar al final del documento
            doc.add_paragraph("")
            doc.add_paragraph(f"Santiago, {datetime.now().strftime('%d de %B de %Y')}")
            doc.add_paragraph("")
            
            # Insertar imagen de firma
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(str(FIRMA_TIMBRE_PATH), width=Inches(2.5))
            
            doc.add_paragraph(f"{EMPRESA['representante_legal']}")
            doc.add_paragraph(f"RUT: {EMPRESA['rut_rep_legal']}")
            doc.add_paragraph(f"{EMPRESA['cargo_rep_legal']}")
            doc.add_paragraph(f"{EMPRESA['razon_social']}")
            doc.add_paragraph(f"RUT Empresa: {EMPRESA['rut']}")
        
        # Guardar
        nombre_salida = Path(archivo_path).stem + "_COMPLETADO.docx"
        ruta_salida = Path(archivo_path).parent / nombre_salida
        doc.save(str(ruta_salida))
        
        print(f"Documento guardado: {ruta_salida}")
        return str(ruta_salida)
        
    except ImportError:
        print("python-docx no instalado. Instalando...")
        os.system("pip install python-docx --break-system-packages -q")
        return llenar_anexo_word(archivo_path, analisis)
    except Exception as e:
        print(f"Error llenando Word: {e}")
        return None


def procesar_anexos(carpeta_licitacion):
    """Procesa todos los anexos de una carpeta de licitación."""
    
    carpeta = Path(carpeta_licitacion)
    if not carpeta.exists():
        print(f"La carpeta {carpeta} no existe.")
        return
    
    archivos = list(carpeta.glob("*.pdf")) + list(carpeta.glob("*.docx")) + list(carpeta.glob("*.doc"))
    
    if not archivos:
        print("No se encontraron anexos (.pdf, .docx, .doc)")
        return
    
    print(f"\nProcesando {len(archivos)} archivos...")
    
    resultados = []
    for archivo in archivos:
        print(f"\n--- {archivo.name} ---")
        print("Analizando con IA...")
        
        analisis = analizar_anexo_con_ia(archivo)
        
        if "error" in analisis:
            print(f"Error: {analisis['error']}")
            continue
        
        print(f"Documento: {analisis.get('titulo_documento', 'Desconocido')}")
        print(f"Campos a llenar: {len(analisis.get('campos_a_llenar', []))}")
        print(f"Requiere firma: {analisis.get('requiere_firma', False)}")
        
        if archivo.suffix.lower() == ".docx":
            ruta_completado = llenar_anexo_word(archivo, analisis)
            if ruta_completado:
                resultados.append({
                    "archivo": archivo.name,
                    "completado": Path(ruta_completado).name,
                    "analisis": analisis
                })
        else:
            print(f"  (PDF - análisis completado, llenar manualmente)")
            resultados.append({
                "archivo": archivo.name,
                "completado": None,
                "analisis": analisis
            })
    
    # Generar checklist
    import json
    checklist_path = carpeta / "CHECKLIST_LICITACION.json"
    with open(checklist_path, "w") as f:
        json.dump(resultados, f, ensure_ascii=False, indent=2)
    
    print(f"\nChecklist guardado: {checklist_path}")
    return resultados


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python preparar_anexos.py <carpeta_con_anexos>")
        print("Ejemplo: python preparar_anexos.py ./licitacion_001")
    else:
        procesar_anexos(sys.argv[1])
